"""
PDF to DOC Converter Component

This module provides functionality to convert PDF files to DOC/DOCX format.
Uses multiple approaches including OCR, text extraction, and formatting preservation.
"""

import os
import sys
import tempfile
from pathlib import Path

try:
    from PyPDF2 import PdfReader
    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import pytesseract
    import pdf2image
    from PIL import Image
except ImportError as e:
    print(f"Warning: Missing dependency for PDF to DOC conversion: {e}")
    print("Please install: pip install PyPDF2 pymupdf python-docx pytesseract pdf2image pillow")

def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text from PDF using multiple methods for better accuracy.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Extracted text content
    """
    text_content = ""
    
    try:
        # Method 1: PyPDF2 text extraction
        with open(pdf_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
        
        # If PyPDF2 didn't extract much text, try PyMuPDF
        if len(text_content.strip()) < 100:
            doc = fitz.open(pdf_path)
            text_content = ""
            for page in doc:
                text_content += page.get_text() + "\n"
            doc.close()
            
    except Exception as e:
        print(f"Text extraction error: {e}")
        text_content = ""
    
    return text_content

def extract_images_from_pdf(pdf_path: str, output_dir: str) -> list:
    """
    Extract images from PDF for inclusion in DOC.
    
    Args:
        pdf_path: Path to the PDF file
        output_dir: Directory to save extracted images
        
    Returns:
        List of extracted image paths
    """
    image_paths = []
    
    try:
        # Convert PDF pages to images
        images = pdf2image.convert_from_path(pdf_path, dpi=200)
        
        for i, image in enumerate(images):
            image_path = os.path.join(output_dir, f"page_{i+1}.png")
            image.save(image_path, "PNG")
            image_paths.append(image_path)
            
    except Exception as e:
        print(f"Image extraction error: {e}")
    
    return image_paths

def create_doc_from_text_and_images(text_content: str, image_paths: list, output_path: str) -> bool:
    """
    Create DOC file from extracted text and images.
    
    Args:
        text_content: Extracted text content
        image_paths: List of image paths
        output_path: Output DOC file path
        
    Returns:
        True if successful, False otherwise
    """
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('Converted PDF Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add text content
        if text_content.strip():
            paragraphs = text_content.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # Add images
        for i, image_path in enumerate(image_paths):
            if os.path.exists(image_path):
                # Add page break before image (except first image)
                if i > 0:
                    doc.add_page_break()
                
                # Add image caption
                doc.add_paragraph(f"Page {i+1}")
                
                # Add image
                try:
                    doc.add_picture(image_path, width=Inches(6))
                except Exception as e:
                    print(f"Error adding image {image_path}: {e}")
        
        # Save document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating DOC file: {e}")
        return False

def pdf_to_doc(pdf_path: str, output_path: str, include_images: bool = True, 
               ocr_enabled: bool = False) -> bool:
    """
    Convert PDF to DOC format.
    
    Args:
        pdf_path: Path to input PDF file
        output_path: Path to output DOC file
        include_images: Whether to include images in the output
        ocr_enabled: Whether to use OCR for text extraction
        
    Returns:
        True if conversion successful, False otherwise
    """
    try:
        # Validate input file
        if not os.path.exists(pdf_path):
            print(f"Error: PDF file not found: {pdf_path}")
            return False
        
        # Create temporary directory for extracted content
        with tempfile.TemporaryDirectory() as temp_dir:
            print(f"Converting PDF to DOC: {pdf_path}")
            
            # Extract text
            text_content = extract_text_from_pdf(pdf_path)
            
            # Use OCR if enabled and little text was extracted
            if ocr_enabled and len(text_content.strip()) < 100:
                print("Using OCR for text extraction...")
                text_content = extract_text_with_ocr(pdf_path)
            
            # Extract images if requested
            image_paths = []
            if include_images:
                print("Extracting images...")
                image_paths = extract_images_from_pdf(pdf_path, temp_dir)
            
            # Create DOC file
            print("Creating DOC file...")
            success = create_doc_from_text_and_images(text_content, image_paths, output_path)
            
            if success:
                print(f"✅ Successfully converted to: {output_path}")
            else:
                print("❌ Failed to create DOC file")
            
            return success
            
    except Exception as e:
        print(f"❌ Conversion error: {e}")
        return False

def extract_text_with_ocr(pdf_path: str) -> str:
    """
    Extract text from PDF using OCR (Optical Character Recognition).
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Extracted text content
    """
    try:
        # Convert PDF to images
        images = pdf2image.convert_from_path(pdf_path, dpi=300)
        
        text_content = ""
        for i, image in enumerate(images):
            print(f"Processing page {i+1} with OCR...")
            
            # Extract text using OCR
            page_text = pytesseract.image_to_string(image, lang='eng')
            text_content += f"Page {i+1}:\n{page_text}\n\n"
        
        return text_content
        
    except Exception as e:
        print(f"OCR extraction error: {e}")
        return ""

def batch_pdf_to_doc(input_dir: str, output_dir: str, include_images: bool = True,
                    ocr_enabled: bool = False) -> dict:
    """
    Convert multiple PDF files to DOC format.
    
    Args:
        input_dir: Directory containing PDF files
        output_dir: Directory to save DOC files
        include_images: Whether to include images in the output
        ocr_enabled: Whether to use OCR for text extraction
        
    Returns:
        Dictionary with conversion results
    """
    results = {
        'total_files': 0,
        'successful': 0,
        'failed': 0,
        'errors': []
    }
    
    try:
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Find all PDF files
        pdf_files = list(Path(input_dir).glob("*.pdf"))
        results['total_files'] = len(pdf_files)
        
        print(f"Found {len(pdf_files)} PDF files to convert")
        
        for pdf_file in pdf_files:
            try:
                # Generate output filename
                output_filename = pdf_file.stem + ".docx"
                output_path = os.path.join(output_dir, output_filename)
                
                print(f"\nConverting: {pdf_file.name}")
                
                # Convert PDF to DOC
                success = pdf_to_doc(
                    str(pdf_file), 
                    output_path, 
                    include_images=include_images,
                    ocr_enabled=ocr_enabled
                )
                
                if success:
                    results['successful'] += 1
                else:
                    results['failed'] += 1
                    results['errors'].append(f"Failed to convert {pdf_file.name}")
                    
            except Exception as e:
                results['failed'] += 1
                error_msg = f"Error converting {pdf_file.name}: {str(e)}"
                results['errors'].append(error_msg)
                print(f"❌ {error_msg}")
        
        print(f"\nBatch conversion completed:")
        print(f"Total: {results['total_files']}")
        print(f"Successful: {results['successful']}")
        print(f"Failed: {results['failed']}")
        
        return results
        
    except Exception as e:
        print(f"❌ Batch conversion error: {e}")
        results['errors'].append(f"Batch conversion error: {str(e)}")
        return results

def main():
    """Command line interface for PDF to DOC conversion."""
    if len(sys.argv) < 3:
        print("Usage: python pdf_to_doc.py <input_pdf> <output_doc> [--no-images] [--ocr]")
        print("Example: python pdf_to_doc.py input.pdf output.docx --ocr")
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    output_path = sys.argv[2]
    
    # Parse options
    include_images = "--no-images" not in sys.argv
    ocr_enabled = "--ocr" in sys.argv
    
    # Convert PDF to DOC
    success = pdf_to_doc(pdf_path, output_path, include_images, ocr_enabled)
    
    if success:
        print("✅ Conversion completed successfully!")
        sys.exit(0)
    else:
        print("❌ Conversion failed!")
        sys.exit(1)

if __name__ == "__main__":
    main() 